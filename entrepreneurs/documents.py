from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from datetime import datetime

def generate_invoice(entrepreneur, bank_details, items, form_data):
    document = Document()
    
    # Настройка стилей
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('СЧЕТ НА ОПЛАТУ № 1')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    document.add_paragraph(f"от {form_data['date'].strftime('%d.%m.%Y')}")
    document.add_paragraph()
    
    # Поставщик
    supplier = document.add_paragraph()
    supplier.add_run('Поставщик: ').bold = True
    supplier.add_run(f"ИП {entrepreneur.company_name}, ИНН {entrepreneur.inn}, ОГРНИП {entrepreneur.ogrnip}")
    
    # Банковские реквизиты
    bank = document.add_paragraph()
    bank.add_run('Банковские реквизиты: ').bold = True
    bank.add_run(f"{bank_details.bank_name}, БИК {bank_details.bik}, ")
    bank.add_run(f"к/с {bank_details.correspondent_account}, р/с {bank_details.payment_account}")
    
    document.add_paragraph()
    
    # Покупатель
    customer = document.add_paragraph()
    customer.add_run('Покупатель: ').bold = True
    customer.add_run(f"{form_data['customer_name']}, ИНН {form_data['customer_inn']}, ")
    customer.add_run(f"адрес: {form_data['customer_address']}")
    
    document.add_paragraph()
    
    # Таблица с товарами
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Кол-во'
    hdr_cells[3].text = 'Ед.'
    hdr_cells[4].text = 'Цена'
    hdr_cells[5].text = 'Сумма'
    
    # Добавление товаров
    total = 0
    for i, item in enumerate(items, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.description
        row_cells[2].text = str(item.quantity)
        row_cells[3].text = item.unit
        row_cells[4].text = f"{item.price:.2f}"
        sum_row = item.quantity * item.price
        row_cells[5].text = f"{sum_row:.2f}"
        total += sum_row
    
    # Итого
    total_row = table.add_row().cells
    for i in range(4):
        total_row[i].text = ''
    total_row[4].text = 'Итого:'
    total_row[5].text = f"{total:.2f}"
    
    document.add_paragraph()
    
    # Подпись
    sign = document.add_paragraph()
    sign.add_run(f"ИП {entrepreneur.company_name}").bold = True
    sign.add_run(' ___________________________')
    

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream